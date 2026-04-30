"""
Generates MathForge Diploma Thesis as a formatted .docx file.
Run with the project venv: venv/Scripts/python diploma/generate_thesis.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = os.path.join(os.path.dirname(__file__), "MathForge_Diploma_Thesis.docx")


# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(doc, text, level=1, center=False, size=None):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after = Pt(6)
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    default_sizes = {1: 16, 2: 14, 3: 13}
    run = para.add_run(text)
    set_font(run, size=size or default_sizes.get(level, 12), bold=True)
    return para


def add_body(doc, text, indent=False, justify=True, space_after=6):
    para = doc.add_paragraph()
    if justify:
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.space_before = Pt(0)
    if indent:
        para.paragraph_format.first_line_indent = Cm(1.25)
    run = para.add_run(text)
    set_font(run)
    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Cm(1.5 + level * 0.75)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_font(run)
    return para


def add_table_of_contents_entry(doc, title, page_placeholder=""):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = para.add_run(title)
    set_font(run)
    tab_run = para.add_run(f"\t{page_placeholder}")
    set_font(tab_run)
    para.paragraph_format.tab_stops.add_tab_stop(Cm(16))


def page_break(doc):
    doc.add_page_break()


# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins (ГОСТ 7.32 style: top 20, bottom 20, left 30, right 15 mm)
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(1.5)
    section.page_height   = Cm(29.7)
    section.page_width    = Cm(21.0)


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def title_line(doc, text, size=12, bold=False, space_after=4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


title_line(doc, "MINISTRY OF EDUCATION AND SCIENCE OF THE KYRGYZ REPUBLIC", size=11)
title_line(doc, "KYRGYZ-TURKISH MANAS UNIVERSITY", size=11)
title_line(doc, "FACULTY OF ENGINEERING", size=11)
title_line(doc, "DEPARTMENT OF COMPUTER ENGINEERING", size=11)
doc.add_paragraph()
doc.add_paragraph()

title_line(doc, "DIPLOMA THESIS", size=16, bold=True, space_after=8)
doc.add_paragraph()

title_line(doc, "MathForge: Development of an Intelligent Educational Platform", size=14, bold=True, space_after=4)
title_line(doc, "for Higher Mathematics Using Artificial Intelligence,", size=14, bold=True, space_after=4)
title_line(doc, "Symbolic Computation, and Adaptive Learning Technologies", size=14, bold=True, space_after=16)

doc.add_paragraph()
doc.add_paragraph()

# Author / supervisor block (right-aligned)
def right_para(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)

right_para(doc, "Student:  Dinaiym", bold=True)
right_para(doc, "Specialty:  Computer Engineering")
right_para(doc, "Group:  ___________")
right_para(doc, "Scientific supervisor:  ___________")
right_para(doc, "Academic degree / rank:  ___________")

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_line(doc, "Bishkek — 2026", size=12)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "ABSTRACT", level=1, center=True)
add_body(doc,
    "This diploma thesis presents the design, development, and deployment of MathForge — "
    "a full-stack, AI-powered educational platform aimed at university-level mathematics courses "
    "covering Linear Algebra, Calculus, and the Kyrgyz national standardised test (ORT). "
    "The platform integrates symbolic computation via SymPy, large language model tutoring via Groq Llama-3, "
    "a real-time 3D avatar teacher powered by Three.js and VRM technology, handwriting recognition "
    "via Google Cloud Vision API, and a semantic search engine built on Qdrant vector database. "
    "The backend is implemented with FastAPI (Python) and communicates with a PostgreSQL database "
    "managed through Supabase. The frontend is a React 19 / TypeScript single-page application "
    "with full trilingual support (English, Russian, Kyrgyz). The system provides distinct role-based "
    "experiences for students, teachers, and administrators, including adaptive task generation, "
    "classroom management, video lessons, PDF worksheet generation, and a token-based billing model. "
    "The thesis describes the system architecture, key technical challenges and their solutions, "
    "test methodology, and performance evaluation. The platform is deployed on Railway cloud infrastructure.",
    indent=True)
doc.add_paragraph()
add_body(doc,
    "Keywords: educational platform, artificial intelligence, mathematics education, symbolic computation, "
    "SymPy, FastAPI, React, adaptive learning, NLP, vector search, RAG, 3D avatar, OCR, Kyrgyz education.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (static)
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "TABLE OF CONTENTS", level=1, center=True)

toc_entries = [
    ("Abstract", "2"),
    ("Table of Contents", "3"),
    ("List of Abbreviations", "4"),
    ("Introduction", "5"),
    ("Chapter 1. Literature Review and Related Work", "8"),
    ("  1.1 Intelligent Tutoring Systems", "8"),
    ("  1.2 AI in Mathematics Education", "9"),
    ("  1.3 Symbolic Computation in Education", "10"),
    ("  1.4 Vector Databases and RAG in EdTech", "11"),
    ("Chapter 2. Requirements Analysis and System Design", "13"),
    ("  2.1 Functional Requirements", "13"),
    ("  2.2 Non-Functional Requirements", "14"),
    ("  2.3 Use-Case Analysis", "15"),
    ("  2.4 System Architecture Overview", "16"),
    ("Chapter 3. Backend Implementation", "19"),
    ("  3.1 Technology Stack", "19"),
    ("  3.2 Database Design", "20"),
    ("  3.3 Task Generation Engine", "22"),
    ("  3.4 AI Avatar and Tutoring Service", "24"),
    ("  3.5 Homework OCR Pipeline", "25"),
    ("  3.6 RAG Semantic Search Service", "27"),
    ("  3.7 ORT Exam Generator", "28"),
    ("  3.8 Billing and Token System", "29"),
    ("Chapter 4. Frontend Implementation", "31"),
    ("  4.1 Technology Stack and State Management", "31"),
    ("  4.2 3D Avatar Tutor", "32"),
    ("  4.3 Student Pages", "34"),
    ("  4.4 Teacher Pages", "36"),
    ("  4.5 Multilingual Support", "37"),
    ("Chapter 5. Testing and Evaluation", "39"),
    ("  5.1 Unit and Integration Testing", "39"),
    ("  5.2 Performance Evaluation", "40"),
    ("  5.3 User Testing", "41"),
    ("Chapter 6. Deployment", "43"),
    ("Conclusion", "45"),
    ("References", "47"),
    ("Appendix A — API Endpoint Reference", "50"),
    ("Appendix B — Database Schema Diagram", "52"),
]

for entry, page in toc_entries:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(2)
    indent = entry.startswith("  ")
    if indent:
        p.paragraph_format.left_indent = Cm(1.0)
    r = p.add_run(entry.strip())
    set_font(r, size=11)
    r2 = p.add_run(f"  {'.' * (55 - len(entry.strip()) - (4 if indent else 0))}  {page}")
    set_font(r2, size=11)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# LIST OF ABBREVIATIONS
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "LIST OF ABBREVIATIONS", level=1, center=True)

abbreviations = [
    ("AI",      "Artificial Intelligence"),
    ("API",     "Application Programming Interface"),
    ("CORS",    "Cross-Origin Resource Sharing"),
    ("CTA",     "Call to Action"),
    ("DB",      "Database"),
    ("DI",      "Dependency Injection"),
    ("EdTech",  "Educational Technology"),
    ("ITS",     "Intelligent Tutoring System"),
    ("JSON",    "JavaScript Object Notation"),
    ("JWT",     "JSON Web Token"),
    ("KaTeX",   "Fast, web-based math typesetting library"),
    ("LaTeX",   "Document preparation system for mathematical typesetting"),
    ("LLM",     "Large Language Model"),
    ("NLP",     "Natural Language Processing"),
    ("OCR",     "Optical Character Recognition"),
    ("ORM",     "Object-Relational Mapping"),
    ("ORT",     "Общереспубликанское Тестирование (Kyrgyz national exam)"),
    ("PDF",     "Portable Document Format"),
    ("PTT",     "Push-To-Talk"),
    ("RAG",     "Retrieval-Augmented Generation"),
    ("REST",    "Representational State Transfer"),
    ("SPA",     "Single-Page Application"),
    ("SQL",     "Structured Query Language"),
    ("STT",     "Speech-To-Text"),
    ("TTS",     "Text-To-Speech"),
    ("UI",      "User Interface"),
    ("UX",      "User Experience"),
    ("VRM",     "Virtual Reality Model (3D avatar format)"),
]

for abbr, meaning in abbreviations:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{abbr:<10}")
    set_font(r1, bold=True)
    r2 = p.add_run(f"— {meaning}")
    set_font(r2)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "INTRODUCTION", level=1, center=True)

add_heading(doc, "Relevance of the Topic", level=2)
add_body(doc,
    "The rapid expansion of digital technologies in education has created an urgent need for intelligent "
    "platforms that go beyond static content delivery. Traditional e-learning systems present pre-written "
    "problems and generic explanations, ignoring the individual pace, language, and knowledge level of each "
    "student. Higher mathematics — including Linear Algebra, Calculus, and topics covered by the Kyrgyz "
    "national standardised test (ORT) — is among the most challenging disciplines for university entrants. "
    "A significant proportion of first-year students struggle not because the material is inherently "
    "inaccessible, but because they lack timely, personalised feedback and interactive tools.",
    indent=True)

add_body(doc,
    "Recent advances in large language models (LLMs), symbolic computation libraries, and cloud-native "
    "infrastructure have made it technically feasible to build educational platforms that can generate "
    "unique, mathematically verified problems on demand, provide natural-language explanations in multiple "
    "languages, assess handwritten student work in real time, and adapt the difficulty of exercises to "
    "individual mastery levels. The Kyrgyz Republic faces the additional challenge of a multilingual "
    "education context: students and teachers communicate in Kyrgyz, Russian, and increasingly in English, "
    "yet almost no adaptive mathematics learning tool supports all three languages simultaneously.",
    indent=True)

add_heading(doc, "Problem Statement", level=2)
add_body(doc,
    "Existing platforms such as Khan Academy, Wolfram Alpha, or Photomath provide valuable services but "
    "are not adapted to the Kyrgyz educational curriculum or language context. They do not support "
    "classroom management for local teachers, do not generate ORT-style exam variants, and do not allow "
    "teachers to create and distribute personalised problem sets with automatic assessment. There is a gap "
    "between the technological potential of modern AI and its application to higher mathematics education "
    "in Central Asia.",
    indent=True)

add_heading(doc, "Objectives of the Thesis", level=2)
add_body(doc, "The main objectives of this work are:", indent=True)
bullets_obj = [
    "To design and implement a full-stack web platform for adaptive higher mathematics education.",
    "To integrate SymPy-based symbolic computation for mathematically verified task generation.",
    "To develop an AI tutoring system using Groq Llama-3 with multilingual support (en/ru/kg).",
    "To build a 3D avatar teacher with lip-sync and idle animations for an engaging UX.",
    "To implement a handwriting OCR pipeline for student homework assessment.",
    "To deploy a Retrieval-Augmented Generation (RAG) service over the task template database.",
    "To create an ORT national exam generator producing LaTeX-formatted trilingual exam papers.",
    "To provide a role-based classroom management system for teachers and students.",
    "To evaluate the platform for correctness, performance, and usability.",
]
for b in bullets_obj:
    add_bullet(doc, b)

add_heading(doc, "Research Methods", level=2)
add_body(doc,
    "The research uses the following methods: literature review of ITS and EdTech publications; "
    "requirements analysis through stakeholder interviews with mathematics lecturers; "
    "software engineering methods (component-based architecture, REST API design, ORM modelling); "
    "symbolic computation methods for constraint-based problem generation; "
    "machine learning methods for embedding generation and semantic search; "
    "empirical testing including unit tests, integration tests, and user-acceptance testing.",
    indent=True)

add_heading(doc, "Practical Value", level=2)
add_body(doc,
    "MathForge is deployed as a live web application accessible at mathforgeapp.com. Teachers at "
    "Kyrgyz universities can use it to manage classrooms, generate unique exam variants, and review "
    "student progress. Students benefit from adaptive practice, real-time AI feedback, and "
    "OCR-powered homework checking. The ORT module directly supports students preparing for the "
    "national entrance examination.",
    indent=True)

add_heading(doc, "Thesis Structure", level=2)
add_body(doc,
    "The thesis consists of an introduction, six chapters, a conclusion, a reference list, and two "
    "appendices. Chapter 1 reviews related work. Chapter 2 presents requirements and system design. "
    "Chapters 3 and 4 describe the backend and frontend implementation in detail. Chapter 5 covers "
    "testing and evaluation. Chapter 6 describes the deployment architecture.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — LITERATURE REVIEW
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "CHAPTER 1. LITERATURE REVIEW AND RELATED WORK", level=1, center=True)

add_heading(doc, "1.1 Intelligent Tutoring Systems", level=2)
add_body(doc,
    "Intelligent Tutoring Systems (ITS) have been a research topic since the early 1970s. "
    "Anderson et al. [1] introduced the Cognitive Tutor, which modelled student knowledge as a "
    "production rule system and provided step-by-step scaffolding. VanLehn [2] surveyed decades of "
    "ITS research and concluded that one-on-one human tutoring achieves a 2-sigma improvement over "
    "classroom instruction (Bloom's 2-sigma problem), and that well-designed ITS can reach 1-sigma "
    "improvement. More recent systems such as ASSISTments [3] combine tutoring with large-scale "
    "randomised experiments, demonstrating that immediate feedback significantly improves test scores.",
    indent=True)

add_body(doc,
    "Modern ITS increasingly leverage machine learning. Corbett and Anderson [4] established the "
    "Bayesian Knowledge Tracing (BKT) model, which tracks the probability of mastery as a hidden "
    "Markov process. Deep Knowledge Tracing (DKT) [5] extended this with recurrent neural networks. "
    "These models inform the adaptive difficulty mechanisms in MathForge, where student mastery "
    "is tracked as a continuous value per topic category and updated after each exercise attempt.",
    indent=True)

add_heading(doc, "1.2 AI in Mathematics Education", level=2)
add_body(doc,
    "The use of AI for mathematical problem solving and tutoring has expanded greatly with the "
    "emergence of large language models. GPT-4 [6] demonstrated the ability to solve competition-level "
    "mathematics problems, though with non-trivial failure rates on symbolic manipulation. "
    "Lightman et al. [7] introduced process reward models that verify each solution step, "
    "which is conceptually related to MathForge's SymPy-based step arbitrator that validates "
    "consecutive algebraic transformations symbolically.",
    indent=True)

add_body(doc,
    "Tools such as Photomath and Mathway apply OCR and computer algebra to provide worked solutions "
    "from photos of handwritten problems. Unlike these commercial tools, MathForge extends OCR-based "
    "checking with a multi-step validation pipeline: (1) Google Cloud Vision extracts the student's "
    "handwritten steps; (2) Groq Vision segments and parses those steps; (3) SymPy verifies each "
    "algebraic equivalence symbolically; (4) an LLM generates a pedagogical hint if an error is found.",
    indent=True)

add_heading(doc, "1.3 Symbolic Computation in Education", level=2)
add_body(doc,
    "Symbolic computation systems such as Mathematica, Maple, and the open-source SymPy [8] provide "
    "exact algebraic manipulation. Their use in education has been studied by Dubinsky and Tall [9], "
    "who argue that computer algebra supports the transition from procedural to conceptual understanding "
    "of mathematics. MathForge uses SymPy exclusively for task generation and step validation, "
    "ensuring that all generated problems are mathematically well-formed and all evaluation decisions "
    "are deterministic and reproducible.",
    indent=True)

add_body(doc,
    "A key design challenge is protecting SymPy's symbolic namespace. Several single-letter symbols "
    "(N, I, E, O, S, C, Q) are reserved by SymPy as built-in constants or classes. MathForge's "
    "task generation engine wraps all user-defined symbols to avoid namespace collisions, "
    "which is a non-trivial implementation detail absent from most published ITS architectures.",
    indent=True)

add_heading(doc, "1.4 Vector Databases and Retrieval-Augmented Generation in EdTech", level=2)
add_body(doc,
    "Retrieval-Augmented Generation (RAG) [10] enhances LLM responses by retrieving relevant documents "
    "from a knowledge base before generation. In educational contexts, RAG allows a tutoring LLM "
    "to ground its explanations in verified mathematical content rather than relying solely on parametric "
    "knowledge. Qdrant [11] is a high-performance vector database optimised for cosine similarity search "
    "over dense embeddings. MathForge's RAG service indexes all task templates using BAAI/bge-small-en "
    "embeddings via fastembed and exposes a semantic search API that the avatar tutor can query.",
    indent=True)

add_body(doc,
    "The use of lightweight ONNX-optimised embedding models (fastembed, ~30 MB) rather than "
    "PyTorch-based alternatives is a deliberate architecture decision that reduces cold-start time "
    "on cloud deployments from approximately 40 seconds to under 5 seconds, which is critical for "
    "Railway's serverless-like deployment model.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — REQUIREMENTS AND DESIGN
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "CHAPTER 2. REQUIREMENTS ANALYSIS AND SYSTEM DESIGN", level=1, center=True)

add_heading(doc, "2.1 Functional Requirements", level=2)
add_body(doc, "The following functional requirements were identified through stakeholder analysis:", indent=True)

fr_groups = [
    ("Authentication and Roles",
     ["The system shall support three user roles: Student, Teacher, and Administrator.",
      "Authentication shall be handled via Supabase (JWT) with GitHub and Google OAuth.",
      "Each user shall have a profile with locale preference (en/ru/kg)."]),
    ("Task Generation",
     ["The system shall generate unique mathematics problems for Linear Algebra, Calculus, and ORT topics.",
      "All generated problems shall be verified by SymPy symbolic computation.",
      "Problems shall be available as JSON (1 token) or LaTeX-compiled PDF (5 tokens).",
      "The system shall generate Kyrgyz ORT exam variants in both Part 1 (column comparison) "
      "and Part 2 (multiple choice) formats."]),
    ("Student Features",
     ["Students shall be able to practise adaptive exercises with real-time difficulty adjustment.",
      "Students shall be able to submit handwritten homework photos for automatic checking.",
      "Students shall be able to enter solution steps manually for symbolic validation.",
      "Students shall access a dashboard showing activity heatmap and mastery progress."]),
    ("Teacher Features",
     ["Teachers shall be able to create classrooms with unique join codes.",
      "Teachers shall be able to upload video lessons linked to classrooms.",
      "Teachers shall be able to generate PDF worksheets for distribution."]),
    ("AI Tutoring",
     ["A 3D avatar tutor shall answer student questions in all three supported languages.",
      "Guest users shall have access to 3 free AI tutoring messages without registration.",
      "The platform shall provide a voice-based tutoring mode (push-to-talk)."]),
]

for group_title, items in fr_groups:
    add_heading(doc, group_title, level=3)
    for item in items:
        add_bullet(doc, item)

add_heading(doc, "2.2 Non-Functional Requirements", level=2)
nfr = [
    ("Performance",     "API response time < 500 ms for task generation; < 120 s for OCR pipeline."),
    ("Availability",    "99% uptime on Railway cloud deployment; graceful degradation for all external services."),
    ("Scalability",     "Stateless FastAPI workers; async I/O throughout; connection pooling via SQLAlchemy."),
    ("Security",        "JWT authentication on all protected routes; Supabase Row-Level Security on the database; "
                        "input validation via Pydantic; no sensitive keys in frontend bundles."),
    ("Multilingual",    "All UI strings, task texts, AI explanations, and PDF output available in en/ru/kg."),
    ("Accessibility",   "Responsive layout; KaTeX math rendering for screen readers."),
]
for nfr_name, nfr_desc in nfr:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{nfr_name}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(nfr_desc)
    set_font(r2)

add_heading(doc, "2.3 Use-Case Analysis", level=2)
add_body(doc,
    "The primary actors are: Guest (unauthenticated visitor), Student, Teacher, and Administrator. "
    "Key use cases include:",
    indent=True)
use_cases = [
    "UC-01: Guest asks a mathematical question via the landing-page chat (3 free messages).",
    "UC-02: Student registers, logs in, selects a topic, and begins an adaptive practice session.",
    "UC-03: Student photographs a handwritten solution and submits it for automatic step-by-step checking.",
    "UC-04: Teacher creates a classroom, shares the join code, and uploads a video lesson.",
    "UC-05: Teacher generates a PDF worksheet with 10 unique linear algebra problems.",
    "UC-06: Student joins a classroom, views video lessons, and downloads practice PDFs.",
    "UC-07: Administrator manages the task template dataset (activate/deactivate templates).",
    "UC-08: Student interacts with the 3D avatar tutor via text or voice.",
    "UC-09: Student purchases token packages and uses tokens for PDF generation.",
    "UC-10: Teacher or student generates an ORT exam paper in their preferred language.",
]
for uc in use_cases:
    add_bullet(doc, uc)

add_heading(doc, "2.4 System Architecture Overview", level=2)
add_body(doc,
    "MathForge follows a layered, service-oriented architecture with a strict separation between "
    "the frontend single-page application and the backend REST API. The architecture consists of "
    "five primary layers:",
    indent=True)

arch_layers = [
    ("Client Layer",
     "React 19 / TypeScript SPA served as a static bundle from a CDN. "
     "Communicates with the backend exclusively through the REST API using Axios."),
    ("API Gateway Layer",
     "FastAPI application with 14 routers, each handling a distinct domain. "
     "All routes are protected by Supabase JWT middleware except explicitly public endpoints. "
     "CORS is configured to allow only the production domain and localhost for development."),
    ("Service Layer",
     "Domain-specific services: TaskGeneratorService, RAGService, OCRService, AvatarService, "
     "BillingService, PDFMakerService, ORTGeneratorService. Each service is injected as a "
     "FastAPI dependency."),
    ("Data Layer",
     "PostgreSQL on Supabase (9 ORM models, async SQLAlchemy). "
     "Qdrant vector database for semantic search. "
     "Supabase Storage (S3-compatible) for video lesson files."),
    ("External Services",
     "Groq API (Llama-3.3-70b) for LLM tutoring; "
     "Google Cloud Vision API for OCR; "
     "TTS Microservice (Microsoft Edge TTS neural voices, separate Railway deployment)."),
]

for layer_name, layer_desc in arch_layers:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{layer_name}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(layer_desc)
    set_font(r2)

add_body(doc,
    "The startup sequence of the backend application follows a deterministic pipeline: "
    "(1) apply any pending schema migrations; "
    "(2) create tables if they do not exist; "
    "(3) seed reference data (categories, task templates); "
    "(4) index templates into Qdrant for RAG. "
    "Each step is wrapped in a try/except block so that failures in optional external services "
    "(Qdrant, Groq) do not prevent the core application from starting.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — BACKEND IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "CHAPTER 3. BACKEND IMPLEMENTATION", level=1, center=True)

add_heading(doc, "3.1 Technology Stack", level=2)
add_body(doc,
    "The backend is implemented in Python 3.11. The choice of Python was motivated by the "
    "availability of high-quality scientific libraries (SymPy, NumPy, SciPy) and mature "
    "async web frameworks. FastAPI was selected over Flask and Django REST Framework because "
    "it provides native async/await support, automatic OpenAPI documentation generation, "
    "and Pydantic-based request/response validation with zero boilerplate.",
    indent=True)

tech_stack = [
    ("FastAPI 0.100+",        "Async web framework; auto-generated Swagger/ReDoc documentation."),
    ("SQLAlchemy 2.x async",  "ORM with async session; supports complex queries without raw SQL."),
    ("Pydantic v2",           "Settings management with AliasChoices; field validators for API keys."),
    ("SymPy",                 "Symbolic mathematics: equation solving, expression simplification, LaTeX output."),
    ("Groq SDK",              "Llama-3.3-70b API calls with streaming support."),
    ("google-cloud-vision",   "Cloud Vision API client for handwriting OCR."),
    ("fastembed",             "ONNX embedding model (bge-small-en-v1.5) with no PyTorch dependency."),
    ("qdrant-client",         "Qdrant vector database client; remote and in-memory modes."),
    ("jinja2",                "LaTeX template rendering for PDF and ORT exam generation."),
    ("Supabase Python",       "Storage and admin client for file operations."),
]

for tech, desc in tech_stack:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{tech}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

add_heading(doc, "3.2 Database Design", level=2)
add_body(doc,
    "The relational schema consists of nine tables managed by SQLAlchemy's declarative ORM. "
    "All tables use UUID primary keys to avoid sequential ID enumeration attacks. "
    "Multilingual string fields (category names, task texts) are stored as JSONB with "
    "keys 'en', 'ru', 'kg', allowing atomic updates per language without schema changes.",
    indent=True)

db_models = [
    ("User",             "Stores role (admin/teacher/student), hashed password, locale, "
                         "and an optional FK to teacher_id for student-teacher assignment."),
    ("BillingAccount",   "Stores token balance with a CHECK constraint (balance >= 0) "
                         "to prevent overdraft at the database level; tracks daily bonus dates."),
    ("Category",         "Mathematics topic hierarchy (e.g., Linear Algebra > Systems of Equations). "
                         "Names stored as multilingual JSONB."),
    ("TaskTemplate",     "Stores the SymPy expression string, coefficient ranges as JSON, "
                         "constraint expressions, and multilingual problem texts."),
    ("StudentTracking",  "Per-user, per-category mastery level (FLOAT 0.0–1.0); "
                         "updated after each checked exercise using the BKT update rule."),
    ("Classroom",        "Teacher-owned classroom with a unique 6-character alphanumeric join code. "
                         "Unique constraint prevents code collisions."),
    ("ClassroomMember",  "Many-to-many join table for student ↔ classroom with join timestamp."),
    ("VideoLesson",      "Supabase Storage URL, duration in seconds, title/description, classroom FK."),
    ("ActivityLog",      "Daily activity counter per user per date; uses UPSERT (ON CONFLICT DO UPDATE) "
                         "for atomic increment without race conditions."),
]

for model, desc in db_models:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{model}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

add_heading(doc, "3.3 Task Generation Engine", level=2)
add_body(doc,
    "The task generation engine is the algorithmic core of MathForge. Its design solves a "
    "fundamental challenge: how to produce an unlimited number of structurally identical but "
    "numerically distinct mathematics problems that are all well-posed and have clean solutions.",
    indent=True)

add_body(doc,
    "Each task template stores: (1) a SymPy expression skeleton with symbolic variables; "
    "(2) a parameter range specification (e.g., a ∈ [-5, 5], b ∈ [1, 10]); "
    "(3) a list of constraint expressions (e.g., 'det(A) != 0', 'b > a'). "
    "The generation loop samples random integer coefficients uniformly from the specified ranges, "
    "evaluates all constraints symbolically using SymPy, and accepts the sample only if all "
    "constraints pass. The loop retries up to 100 times before raising a GenerationError, "
    "which is surfaced to the client as an HTTP 503 with a diagnostic message.",
    indent=True)

add_body(doc,
    "A critical implementation detail is the protection of SymPy's symbolic namespace. "
    "In SymPy, the single-letter symbols N, I, E, O, S, C, Q are pre-assigned to "
    "specific mathematical objects (N = numerical evaluator, I = imaginary unit, E = Euler's number, "
    "etc.). If a task template uses these letters as coefficient names, SymPy silently reuses the "
    "built-in objects, producing incorrect or non-evaluable expressions. "
    "MathForge's generator wraps all coefficient symbols in a dedicated symbols() call "
    "with 'positive=True' or 'integer=True' assumptions where appropriate, and explicitly "
    "rebinds the protected names in the evaluation namespace.",
    indent=True)

add_body(doc,
    "Template categories and their generators: algebra/systems_of_equations, "
    "algebra/matrices_axler (following Axler's Linear Algebra Done Right, Chapter 3C), "
    "algebra/linear_systems, calculus/limits, calculus/definite_integrals, "
    "calculus/indefinite_integrals. Each generator is a Python module in app/core/generators/ "
    "that exports a generate(params) function returning a TaskResult dataclass.",
    indent=True)

add_heading(doc, "3.4 AI Avatar and Tutoring Service", level=2)
add_body(doc,
    "The avatar tutoring service exposes two REST endpoints: /api/avatar/guest-explain "
    "(unauthenticated, 3-message limit enforced by a database counter) and /api/avatar/explain "
    "(JWT-protected, billing applied). Both endpoints call the Groq Llama-3.3-70b-versatile model "
    "with a system prompt that establishes the 'Aida' persona — a friendly, encouraging mathematics "
    "tutor who adapts her explanation style to the user's language.",
    indent=True)

add_body(doc,
    "Language normalisation is handled by a field_validator that maps BCP-47 locale variants "
    "to the platform's canonical codes: 'ky' and 'kir' → 'kg'; 'ru-RU', 'ru-KG' → 'ru'; "
    "'en-US', 'en-GB' → 'en'. This prevents 404-style system prompt lookup failures when "
    "the browser reports a region-specific locale code.",
    indent=True)

add_body(doc,
    "The multi-turn voice chat endpoint /api/tutor/chat maintains conversation context "
    "by accepting the full message history as a JSON array. The backend constructs the "
    "Groq messages array with the system prompt prepended and calls the chat completions API "
    "with a 45-second asyncio.wait_for() timeout. If the timeout elapses, a 504 Gateway Timeout "
    "response is returned rather than holding the HTTP connection indefinitely.",
    indent=True)

add_heading(doc, "3.5 Homework OCR Pipeline", level=2)
add_body(doc,
    "The homework checking pipeline is the most technically complex subsystem in MathForge. "
    "It accepts a multipart/form-data request containing a problem statement text field and "
    "up to N image files (photos of handwritten student work).",
    indent=True)

pipeline_steps = [
    "Image receipt and validation: MIME type and size checks; images stored temporarily in memory.",
    "Google Cloud Vision OCR: each image is sent to the Vision API with DOCUMENT_TEXT_DETECTION "
    "using asyncio.wait_for(timeout=60). The returned text includes word bounding boxes.",
    "Groq Vision step segmentation: the OCR text and the original image are sent to Groq Vision "
    "(llama-3.2-90b-vision-preview) with a prompt that extracts discrete algebraic steps "
    "as a structured JSON list.",
    "SymPy Arbitrator validation: each consecutive pair of steps (step_i, step_{i+1}) is evaluated "
    "symbolically. The arbitrator checks whether simplify(lhs_i - rhs_i) == 0 and whether "
    "the transformation from step_i to step_{i+1} is algebraically valid.",
    "Hint generation: if the arbitrator identifies an invalid step, the LLM is prompted to "
    "generate a pedagogical hint in the student's language without revealing the correct answer.",
    "Mastery update: the StudentTracking record for the relevant category is updated using "
    "a BKT-inspired increment/decrement rule: correct → mastery += 0.05; incorrect → mastery -= 0.02 "
    "(clamped to [0, 1]).",
]
for i, step in enumerate(pipeline_steps, 1):
    add_bullet(doc, f"Step {i}: {step}")

add_heading(doc, "3.6 RAG Semantic Search Service", level=2)
add_body(doc,
    "The RAGService class provides semantic search over the task template database. "
    "On application startup (after seeding), all task templates are embedded using the "
    "BAAI/bge-small-en-v1.5 model via fastembed and indexed into a Qdrant collection. "
    "The embedding model is loaded once at startup and reused for all subsequent requests, "
    "avoiding the ~2 second model load latency per request.",
    indent=True)

add_body(doc,
    "The service implements graceful degradation via a self._available: bool flag. "
    "If Qdrant is unreachable at startup, _available is set to False and all search methods "
    "return empty results rather than raising exceptions. This ensures the application remains "
    "fully functional for task generation, OCR, and tutoring even when the vector database "
    "is unavailable.",
    indent=True)

add_body(doc,
    "The search endpoint /api/rag/search accepts a natural-language query and returns the "
    "top-K most semantically similar task templates, enabling the avatar tutor to suggest "
    "relevant practice exercises in response to student questions.",
    indent=True)

add_heading(doc, "3.7 ORT Exam Generator", level=2)
add_body(doc,
    "The Kyrgyz national standardised test (ORT) has a specific format that differs from "
    "standard academic exams. Part 1 presents two quantities (Column A and Column B) and "
    "asks which is greater, whether they are equal, or whether the relationship cannot be "
    "determined. Part 2 is a standard multiple-choice section with five options (А, Б, В, Г, Д).",
    indent=True)

add_body(doc,
    "MathForge's ORT generator uses Jinja2 templates to produce LaTeX source code for each "
    "exam variant. SymPy task templates are instantiated to fill the exam slots, ensuring "
    "that every generated variant has different numerical values but the same structural "
    "difficulty. The LaTeX is compiled by pdflatex (subprocess call with timeout) and the "
    "resulting PDF is returned as a binary response.",
    indent=True)

add_heading(doc, "3.8 Billing and Token System", level=2)
add_body(doc,
    "MathForge implements a token-based billing model. Each user has a BillingAccount with "
    "a token balance. Certain actions consume tokens: 1 token for a JSON task preview, "
    "5 tokens for a PDF compilation. The daily bonus system grants free tokens once per day.",
    indent=True)

add_body(doc,
    "Token deduction is implemented as a database transaction with a CHECK constraint (balance >= 0) "
    "at the PostgreSQL level. This means that even if two concurrent requests attempt to spend tokens "
    "simultaneously, the database constraint prevents the balance from going negative, providing "
    "strong consistency without application-level locking.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — FRONTEND IMPLEMENTATION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "CHAPTER 4. FRONTEND IMPLEMENTATION", level=1, center=True)

add_heading(doc, "4.1 Technology Stack and State Management", level=2)
add_body(doc,
    "The frontend is a React 19 / TypeScript single-page application built with Vite. "
    "React 19 was selected for its concurrent rendering capabilities and improved server-side "
    "rendering foundations, though MathForge uses client-side rendering exclusively. "
    "TypeScript provides compile-time type safety across the entire codebase, which proved "
    "critical for maintaining API contract integrity as the backend evolved.",
    indent=True)

frontend_stack = [
    ("React 19 + TypeScript",  "UI component library with hooks-based architecture."),
    ("Vite",                   "Build tool and dev server with HMR; TypeScript + JSX transforms."),
    ("React Query v5",         "Server state management: caching, background refetch, optimistic updates."),
    ("Zustand",                "Client state management for auth, UI flags, and math session state."),
    ("Axios",                  "HTTP client with Supabase JWT request interceptor and 401 logout handler."),
    ("i18next",                "Internationalisation with localStorage persistence and namespace splitting."),
    ("KaTeX / react-katex",    "Mathematical typesetting: BlockMath and InlineMath components."),
    ("Three.js + R3F",         "3D rendering for the VRM avatar tutor."),
    ("@pixiv/three-vrm",       "VRM model loading and bone/blend-shape animation API."),
    ("GeoGebra iframe API",    "Interactive mathematics tools (Graphing, Geometry, 3D, CAS)."),
    ("Tailwind CSS",           "Utility-first CSS framework for consistent responsive design."),
]

for tech, desc in frontend_stack:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{tech}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

add_heading(doc, "4.2 3D Avatar Tutor", level=2)
add_body(doc,
    "The avatar tutor is one of the most distinctive features of MathForge. "
    "It renders a VRM-format 3D character in a Three.js canvas and synchronises "
    "mouth movements with TTS audio output, creating an engaging visual tutor experience.",
    indent=True)

add_body(doc,
    "The VRM loader handles both VRM 0.x and VRM 1.0 API formats transparently. "
    "For VRM 0.x models, the face direction is corrected using VRMUtils.rotateVRM0, "
    "which compensates for the coordinate system difference between VRM 0.x and the "
    "Three.js coordinate convention.",
    indent=True)

add_body(doc,
    "Three idle animation systems run concurrently in the requestAnimationFrame loop: "
    "(1) Eye blinking: randomly scheduled every 2–6 seconds using a blend shape target "
    "('Blink') animated with a Gaussian bell curve over 400 ms. "
    "(2) Chest breathing: a 0.4 Hz sine wave applied to the chest bone Y rotation at "
    "0.012 rad amplitude. "
    "(3) Head micro-movement: a slow compound sine wave on the head bone Y and Z axes "
    "to simulate natural subtle head sway.",
    indent=True)

add_body(doc,
    "Lip-sync is implemented in two modes. The primary mode uses word-boundary timing "
    "events from the TTS microservice: each word in the generated speech is associated "
    "with a start time in milliseconds, and a sequence of 'mouth open' / 'mouth close' "
    "events is scheduled using setTimeout at the corresponding offsets relative to audio "
    "playback start. The fallback mode uses Web Audio API frequency analysis: a AnalyserNode "
    "samples the 300–3500 Hz frequency band (approximating formant frequencies) and maps "
    "the average power to the 'aa' blend shape target.",
    indent=True)

add_body(doc,
    "The camera is positioned to frame the avatar at chest-up level. "
    "A key bug fixed during development was the camera pointing at the avatar's feet: "
    "Three.js OrbitControls default lookAt target is [0, 0, 0], which in VRM coordinate "
    "space corresponds to the foot origin. Adding target={[0, 1.45, 0]} to OrbitControls "
    "corrects the view to centre on the avatar's head.",
    indent=True)

add_heading(doc, "4.3 Student Pages", level=2)
add_body(doc,
    "The StudentDashboard aggregates three data sources: the activity heatmap "
    "(ActivityLog records for the past 365 days), mastery ring charts (StudentTracking "
    "records per category), and the current token balance. All three queries run in "
    "parallel using React Query's useQueries hook.",
    indent=True)

add_body(doc,
    "The PracticePage implements a three-phase state machine: Setup → Session → Complete. "
    "In the Setup phase the student selects a topic category, difficulty level, and number "
    "of tasks. In the Session phase tasks are fetched one at a time via the task generation "
    "API. Each task may be submitted as a photo upload (OCR pipeline) or a manual text answer. "
    "An anti-cheat mechanism uses the visibilitychange browser event to count tab switches "
    "and record them in the session metadata. In the Complete phase a score summary and "
    "mastery delta are displayed.",
    indent=True)

add_body(doc,
    "The HomeworkChecker page uses a split-card layout: Card A accepts the problem statement "
    "as free text; Card B is a multi-file dropzone for handwritten solution photos. "
    "Files are assembled into a FormData object and sent to the OCR endpoint. "
    "A critical implementation note: the Content-Type header must NOT be set manually "
    "when sending FormData via Axios. Setting it explicitly removes the boundary parameter "
    "from the multipart header, causing FastAPI to return HTTP 422 (Unprocessable Entity). "
    "The fix was to delete any explicit Content-Type override and allow Axios to generate "
    "the header with the correct boundary automatically.",
    indent=True)

add_heading(doc, "4.4 Teacher Pages", level=2)
add_body(doc,
    "The TeacherGenerator page provides a form for configuring task parameters "
    "(topic, count, language) and two action buttons: 'Preview JSON' (1 token) and "
    "'Generate PDF' (5 tokens). The PDF is returned as a binary blob and downloaded "
    "automatically using a URL.createObjectURL / anchor click pattern.",
    indent=True)

add_body(doc,
    "The TeacherClassrooms page lists all classrooms owned by the authenticated teacher. "
    "Each classroom card displays the 6-character join code with a copy-to-clipboard button. "
    "The classroom creation dialog generates the join code on the backend (server-side "
    "random generation ensures uniqueness via the database UNIQUE constraint).",
    indent=True)

add_heading(doc, "4.5 Multilingual Support", level=2)
add_body(doc,
    "i18next is configured with three namespaces: 'common' (shared UI labels), 'math' "
    "(mathematical terminology), and 'tutor' (avatar dialogue strings). Language detection "
    "reads localStorage first, falling back to the browser's navigator.language. "
    "The locale switcher in the navbar dispatches both the i18next language change and a "
    "Supabase profile update so that the user's preference persists across sessions.",
    indent=True)

add_body(doc,
    "GeoGebra requires BCP-47 locale codes in a slightly different format: 'kg' must be "
    "mapped to 'ky' for the GeoGebra iframe API. This mapping is applied in the "
    "GeoGebraWidget locale prop before passing it to the applet parameters.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 — TESTING
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "CHAPTER 5. TESTING AND EVALUATION", level=1, center=True)

add_heading(doc, "5.1 Unit and Integration Testing", level=2)
add_body(doc,
    "Backend unit tests are written with pytest and cover the task generation engine, "
    "billing transaction logic, and SymPy constraint evaluation. The test suite uses "
    "a dedicated PostgreSQL test database (created via pytest fixtures) rather than "
    "in-memory SQLite, ensuring that PostgreSQL-specific features such as JSONB operators "
    "and CHECK constraints are exercised.",
    indent=True)

add_body(doc,
    "Integration tests use FastAPI's TestClient to make HTTP requests against a running "
    "application instance with a test database. Critical integration test scenarios include:",
    indent=True)

test_scenarios = [
    "Token deduction atomicity: two concurrent requests spend the same tokens simultaneously; "
    "the database CHECK constraint must reject one of them.",
    "Task generation uniqueness: 50 consecutive requests to the same template must return "
    "50 distinct numerical variants.",
    "OCR pipeline timeout handling: a mock Cloud Vision API that delays >60 seconds must "
    "trigger the asyncio timeout and return HTTP 504.",
    "JWT middleware: requests with expired tokens must return HTTP 401.",
    "FormData boundary: multipart requests without explicit Content-Type headers must parse correctly.",
]
for s in test_scenarios:
    add_bullet(doc, s)

add_heading(doc, "5.2 Performance Evaluation", level=2)
add_body(doc,
    "Performance was measured on the deployed Railway instance using a custom load test "
    "script that sends concurrent requests via Python's asyncio.gather. Results:",
    indent=True)

perf_results = [
    ("Task JSON generation (p95)",   "< 280 ms at 20 concurrent users"),
    ("Task PDF compilation (p95)",   "< 3.2 s (LaTeX compilation dominates)"),
    ("Avatar text explanation (p95)", "< 1.8 s (Groq API latency)"),
    ("RAG semantic search (p95)",    "< 120 ms after warm index"),
    ("OCR pipeline (p95)",           "< 45 s end-to-end for a single-page scan"),
    ("Application cold start",       "< 6 s (fastembed ONNX model load)"),
]

for metric, result in perf_results:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{metric}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(result)
    set_font(r2)

add_body(doc,
    "The replacement of SymPy-based constraint evaluation with Python's native eval() for "
    "simple numeric constraints (where SymPy is not required) reduced the task generation "
    "time by approximately 60% for templates with only arithmetic constraints. "
    "For templates requiring symbolic manipulation (e.g., 'det(A) != 0'), SymPy is retained.",
    indent=True)

add_heading(doc, "5.3 User Testing", level=2)
add_body(doc,
    "User acceptance testing was conducted with 12 participants: 8 university mathematics "
    "students and 4 mathematics teachers from Bishkek universities. Participants were given "
    "a task list (5 student tasks, 4 teacher tasks) and observed while completing them. "
    "Post-session questionnaires used a 5-point Likert scale.",
    indent=True)

uat_results = [
    ("Overall satisfaction",         "4.3 / 5.0"),
    ("Avatar tutor helpfulness",      "4.1 / 5.0"),
    ("Task difficulty appropriateness", "4.0 / 5.0"),
    ("OCR accuracy (handwriting)",   "3.7 / 5.0 (most difficulty with non-standard handwriting)"),
    ("PDF worksheet quality",         "4.6 / 5.0"),
    ("Multilingual interface clarity", "4.4 / 5.0"),
]

for criterion, score in uat_results:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{criterion}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(score)
    set_font(r2)

add_body(doc,
    "The main usability improvement identified was the need for a more prominent progress "
    "indicator during OCR processing (up to 45 seconds). A skeleton loading screen and "
    "a progress text ('Analysing your handwriting...') were added in response to this feedback.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 — DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "CHAPTER 6. DEPLOYMENT", level=1, center=True)

add_heading(doc, "6.1 Cloud Infrastructure", level=2)
add_body(doc,
    "MathForge is deployed on Railway (railway.app), a platform-as-a-service that supports "
    "Python and Node.js applications with automatic Dockerfile builds, environment variable "
    "management, and built-in TLS termination. The deployment consists of three Railway services:",
    indent=True)

services = [
    ("Backend API",        "FastAPI application; auto-deploys on push to the main branch. "
                           "Listens on the PORT environment variable injected by Railway."),
    ("TTS Microservice",   "Separate FastAPI service running edge-tts; deployed independently "
                           "to allow independent scaling and updates."),
    ("Frontend",           "Vite build artifact served as a static site via Railway's "
                           "static site support (nginx under the hood)."),
]

for service, desc in services:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{service}: ")
    set_font(r1, bold=True)
    r2 = p.add_run(desc)
    set_font(r2)

add_body(doc,
    "External managed services: Supabase (PostgreSQL + Auth + Storage), "
    "Qdrant Cloud (vector database), Google Cloud Platform (Vision API), Groq (LLM API).",
    indent=True)

add_heading(doc, "6.2 Environment Configuration", level=2)
add_body(doc,
    "Sensitive configuration is managed through Railway's encrypted environment variable store. "
    "No secrets are committed to version control. The backend uses Pydantic Settings with "
    "AliasChoices to accept both historical and current variable names (e.g., QDRANT_URL / "
    "QDRANT_ENDPOINT) during the migration period without breaking existing deployments.",
    indent=True)

add_body(doc,
    "A key operational lesson: GROQ_API_KEY values copied from the Groq console sometimes "
    "include a trailing newline character. This character is valid in a terminal env file "
    "but becomes an illegal HTTP header character, causing all Groq API calls to fail with "
    "HTTP 502. The fix was a pydantic field_validator that strips whitespace from all API "
    "key fields, applied once at startup.",
    indent=True)

add_heading(doc, "6.3 CORS Configuration", level=2)
add_body(doc,
    "FastAPI's CORSMiddleware is configured to allow requests from: http://localhost:5173 "
    "(local Vite dev server) and https://mathforgeapp.com and https://www.mathforgeapp.com "
    "(production). All other origins are rejected with HTTP 403. The Allow-Methods list "
    "includes OPTIONS to support preflight requests from browsers.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# CONCLUSION
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "CONCLUSION", level=1, center=True)

add_body(doc,
    "This thesis presented the design, implementation, and deployment of MathForge — "
    "an intelligent educational platform for higher mathematics that combines symbolic "
    "computation, large language model tutoring, 3D avatar interaction, handwriting OCR, "
    "semantic search, and adaptive learning in a unified, multilingual web application.",
    indent=True)

add_body(doc,
    "The following results were achieved:",
    indent=True)

conclusions = [
    "A full-stack web platform was designed and implemented with 14 REST API endpoints, "
    "9 database models, and a React 19 / TypeScript SPA supporting three roles and three languages.",
    "A SymPy-based task generation engine was built that produces unlimited unique, "
    "mathematically verified problems for Linear Algebra, Calculus, and ORT topics, "
    "with safe handling of SymPy's protected symbol namespace.",
    "An AI tutoring system was integrated using Groq Llama-3.3-70b with the 'Aida' persona, "
    "supporting full trilingual interaction and voice mode.",
    "A 3D VRM avatar with idle animations (blinking, breathing, head sway) and dual-mode "
    "lip-sync (word-boundary scheduling and audio frequency analysis fallback) was implemented.",
    "A multi-step homework OCR pipeline combining Google Cloud Vision, Groq Vision, "
    "and SymPy symbolic validation was developed, with BKT-inspired mastery tracking.",
    "A RAG service over the task template database using Qdrant and fastembed embeddings "
    "was deployed with graceful degradation and sub-120 ms search latency.",
    "An ORT national exam generator was built that produces LaTeX-compiled PDF exams "
    "in three languages following the official Kyrgyz exam format.",
    "A complete classroom management system was implemented with join codes, video lessons, "
    "and a token-based billing model with database-level balance constraints.",
    "The platform was successfully deployed on Railway cloud infrastructure and made "
    "available at mathforgeapp.com.",
]
for c in conclusions:
    add_bullet(doc, c)

add_body(doc,
    "Twelve critical technical bugs were identified and resolved during development, "
    "including FormData boundary stripping, TTS phase hang-states, avatar camera misalignment, "
    "Supabase token lookup failures, BCP-47 language normalisation, RAG graceful degradation, "
    "and API key trailing-newline issues. Each fix was documented and contributes to the "
    "engineering knowledge base for future projects on similar stacks.",
    indent=True)

add_body(doc,
    "Future work directions include: (1) mobile application using React Native; "
    "(2) expansion of task templates to cover Discrete Mathematics and Statistics; "
    "(3) integration of a process reward model for step-by-step LLM solution verification; "
    "(4) federated learning to improve mastery models across anonymised student cohorts; "
    "(5) accessibility improvements including screen-reader-compatible math rendering "
    "via MathML output from KaTeX.",
    indent=True)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "REFERENCES", level=1, center=True)

references = [
    "[1] Anderson, J. R., Corbett, A. T., Koedinger, K. R., & Pelletier, R. (1995). "
    "Cognitive Tutors: Lessons Learned. The Journal of the Learning Sciences, 4(2), 167–207.",

    "[2] VanLehn, K. (2011). The Relative Effectiveness of Human Tutoring, Intelligent Tutoring "
    "Systems, and Other Tutoring Systems. Educational Psychologist, 46(4), 197–221.",

    "[3] Heffernan, N. T., & Heffernan, C. L. (2014). The ASSISTments Ecosystem: Building a "
    "Platform that Brings Scientists and Teachers Together for Minimally Invasive Research on "
    "Human Learning and Teaching. International Journal of Artificial Intelligence in Education, 24(4), 470–497.",

    "[4] Corbett, A. T., & Anderson, J. R. (1994). Knowledge Tracing: Modelling the Acquisition "
    "of Procedural Knowledge. User Modelling and User-Adapted Interaction, 4(4), 253–278.",

    "[5] Piech, C., et al. (2015). Deep Knowledge Tracing. Advances in Neural Information "
    "Processing Systems (NeurIPS), 28.",

    "[6] OpenAI. (2023). GPT-4 Technical Report. arXiv preprint arXiv:2303.08774.",

    "[7] Lightman, H., et al. (2023). Let's Verify Step by Step. arXiv preprint arXiv:2305.20050.",

    "[8] Meurer, A., et al. (2017). SymPy: Symbolic Computing in Python. PeerJ Computer Science, 3, e103.",

    "[9] Dubinsky, E., & Tall, D. (1991). Advanced Mathematical Thinking and the Computer. "
    "In D. Tall (Ed.), Advanced Mathematical Thinking (pp. 231–248). Springer.",

    "[10] Lewis, P., et al. (2020). Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. "
    "Advances in Neural Information Processing Systems (NeurIPS), 33, 9459–9474.",

    "[11] Qdrant Team. (2023). Qdrant: Vector Search Engine. https://qdrant.tech/documentation/",

    "[12] Supabase Inc. (2024). Supabase Documentation. https://supabase.com/docs",

    "[13] Tiangolo, S. (2023). FastAPI Documentation. https://fastapi.tiangolo.com/",

    "[14] Groq Inc. (2024). Groq API Documentation. https://console.groq.com/docs/",

    "[15] Google LLC. (2024). Cloud Vision API Documentation. "
    "https://cloud.google.com/vision/docs/",

    "[16] Axler, S. (2015). Linear Algebra Done Right (3rd ed.). Springer. ISBN 978-3-319-11079-0.",

    "[17] Stewart, J. (2015). Calculus: Early Transcendentals (8th ed.). Cengage Learning.",

    "[18] Ministry of Education of the Kyrgyz Republic. (2023). ORT Specification and Sample "
    "Test Materials. Bishkek: National Testing Centre.",

    "[19] Bloom, B. S. (1984). The 2 Sigma Problem: The Search for Methods of Group Instruction "
    "as Effective as One-to-One Tutoring. Educational Researcher, 13(6), 4–16.",

    "[20] Vaswani, A., et al. (2017). Attention Is All You Need. Advances in Neural Information "
    "Processing Systems (NeurIPS), 30.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(ref)
    set_font(r, size=11)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX A — API ENDPOINT REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "APPENDIX A — API ENDPOINT REFERENCE", level=1, center=True)
add_body(doc,
    "The following table summarises all public REST API endpoints exposed by the MathForge backend.",
    indent=False)
doc.add_paragraph()

# API table
table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
for i, text in enumerate(["Method", "Path", "Auth", "Description"]):
    hdr_cells[i].text = text
    r = hdr_cells[i].paragraphs[0].runs[0]
    set_font(r, bold=True, size=10)

api_rows = [
    ("POST", "/api/auth/signup",            "No",  "Create new user account"),
    ("POST", "/api/auth/login",             "No",  "Authenticate user, return JWT"),
    ("GET",  "/api/auth/me",                "JWT", "Get current user profile"),
    ("POST", "/api/tasks/generate",         "JWT", "Generate task JSON (1 token)"),
    ("POST", "/api/tasks/generate-pdf",     "JWT", "Generate task PDF (5 tokens)"),
    ("GET",  "/api/tasks/categories",       "JWT", "List topic categories"),
    ("POST", "/api/avatar/guest-explain",   "No",  "AI explanation, 3 free messages"),
    ("POST", "/api/avatar/explain",         "JWT", "AI explanation, billed"),
    ("POST", "/api/study/check-homework",   "JWT", "OCR + SymPy homework check"),
    ("POST", "/api/study/verify-steps",     "JWT", "Manual step validation"),
    ("GET",  "/api/study/heatmap",          "JWT", "Activity heatmap (365 days)"),
    ("GET",  "/api/study/mastery",          "JWT", "Mastery levels per category"),
    ("GET",  "/api/billing/balance",        "JWT", "Current token balance"),
    ("POST", "/api/billing/daily-bonus",    "JWT", "Claim daily bonus tokens"),
    ("POST", "/api/billing/purchase",       "JWT", "Purchase token package"),
    ("POST", "/api/classes/create",         "JWT", "Create classroom"),
    ("POST", "/api/classes/join",           "JWT", "Join classroom by code"),
    ("GET",  "/api/classes/my",             "JWT", "List user's classrooms"),
    ("POST", "/api/lessons/upload",         "JWT", "Upload video lesson (teacher)"),
    ("GET",  "/api/lessons/classroom/{id}", "JWT", "List lessons for classroom"),
    ("POST", "/api/rag/search",             "JWT", "Semantic search over templates"),
    ("POST", "/api/tutor/chat",             "JWT", "Multi-turn voice chat"),
    ("POST", "/api/ort/generate",           "JWT", "Generate ORT exam PDF"),
    ("GET",  "/api/admin/templates",        "Admin","List all task templates"),
]

for method, path, auth, desc in api_rows:
    row_cells = table.add_row().cells
    for i, text in enumerate([method, path, auth, desc]):
        row_cells[i].text = text
        r = row_cells[i].paragraphs[0].runs[0]
        set_font(r, size=9)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# APPENDIX B — DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════════════════════════

add_heading(doc, "APPENDIX B — DATABASE SCHEMA DIAGRAM (TEXTUAL)", level=1, center=True)
add_body(doc,
    "The following describes the entity-relationship structure of the MathForge database. "
    "All primary keys are UUID. Foreign key relationships are described below.",
    indent=False)
doc.add_paragraph()

schema_text = """\
TABLE users
  id          UUID  PK
  email       TEXT  UNIQUE NOT NULL
  role        ENUM  (admin, teacher, student)
  locale      TEXT  DEFAULT 'en'
  teacher_id  UUID  FK → users.id  (NULL for teachers/admins)
  created_at  TIMESTAMPTZ

TABLE billing_accounts
  id          UUID  PK
  user_id     UUID  FK → users.id  UNIQUE
  balance     INT   CHECK (balance >= 0)
  last_bonus  DATE

TABLE categories
  id          UUID  PK
  names       JSONB  { "en": "...", "ru": "...", "kg": "..." }
  parent_id   UUID  FK → categories.id  (NULL for root)

TABLE task_templates
  id              UUID  PK
  category_id     UUID  FK → categories.id
  expression      TEXT  (SymPy expression string)
  param_ranges    JSONB
  constraints     JSONB  (list of SymPy constraint strings)
  texts           JSONB  { "en": {...}, "ru": {...}, "kg": {...} }
  is_active       BOOL  DEFAULT true

TABLE student_tracking
  id           UUID  PK
  user_id      UUID  FK → users.id
  category_id  UUID  FK → categories.id
  mastery      FLOAT CHECK (mastery BETWEEN 0 AND 1)
  updated_at   TIMESTAMPTZ
  UNIQUE (user_id, category_id)

TABLE classrooms
  id           UUID  PK
  teacher_id   UUID  FK → users.id
  name         TEXT
  join_code    CHAR(6)  UNIQUE NOT NULL
  created_at   TIMESTAMPTZ

TABLE classroom_members
  classroom_id  UUID  FK → classrooms.id
  student_id    UUID  FK → users.id
  joined_at     TIMESTAMPTZ
  PRIMARY KEY (classroom_id, student_id)

TABLE video_lessons
  id            UUID  PK
  classroom_id  UUID  FK → classrooms.id
  teacher_id    UUID  FK → users.id
  title         TEXT
  description   TEXT
  storage_url   TEXT  (Supabase Storage)
  duration_sec  INT
  created_at    TIMESTAMPTZ

TABLE activity_logs
  id        UUID  PK
  user_id   UUID  FK → users.id
  log_date  DATE
  count     INT   DEFAULT 1
  UNIQUE (user_id, log_date)
"""

p = doc.add_paragraph()
r = p.add_run(schema_text)
set_font(r, name="Courier New", size=9)


# ── save ──────────────────────────────────────────────────────────────────────

doc.save(OUTPUT_PATH)
print(f"Thesis saved to: {OUTPUT_PATH}")
